"""Local text extraction for PDF, DOCX, and plain-text documents."""

import io
import os
from pathlib import Path

import fitz

from app.utils.logger import get_logger

logger = get_logger(__name__)

TEXT_EXTENSIONS = {".txt", ".text", ".md", ".csv"}
WORD_EXTENSIONS = {".docx"}
PDF_CONTENT_TYPES = {"application/pdf", "application/x-pdf"}
TEXT_CONTENT_TYPES = {"text/plain", "text/markdown", "text/csv"}
WORD_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def get_document_type(content, filename=None, content_type=None):
    """Return the supported document type, or ``None`` for unsupported files."""
    normalized_type = (content_type or "").split(";", 1)[0].lower()
    extension = Path(filename or "").suffix.lower()

    if content.startswith(b"%PDF-") or normalized_type in PDF_CONTENT_TYPES or extension == ".pdf":
        return "pdf"
    if normalized_type in WORD_CONTENT_TYPES or extension in WORD_EXTENSIONS:
        return "word"
    if normalized_type in TEXT_CONTENT_TYPES or extension in TEXT_EXTENSIONS:
        return "text"
    return None


def is_supported_document(content, filename=None, content_type=None):
    """Return whether the upload is a supported PDF, DOCX, or text document."""
    return get_document_type(content, filename, content_type) is not None


def _get_tesseract():
    try:
        import pytesseract
    except ImportError as error:
        raise RuntimeError(
            "OCR dependencies are unavailable. Install 'pytesseract' and 'Pillow', "
            "then install the Tesseract OCR application."
        ) from error

    tesseract_cmd = os.getenv("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    return pytesseract


def _ocr_pdf_page(page, dpi):
    """Render a scanned PDF page and extract its text with local Tesseract OCR."""
    from PIL import Image

    pixmap = page.get_pixmap(matrix=fitz.Matrix(dpi / 72, dpi / 72))
    with Image.open(io.BytesIO(pixmap.tobytes("png"))) as image:
        try:
            return _get_tesseract().image_to_string(
                image, lang=os.getenv("OCR_LANGUAGE", "eng")
            )
        except Exception as error:
            raise RuntimeError(f"OCR failed: {error}") from error


def _parse_pdf(content):
    dpi = int(os.getenv("OCR_DPI", "200"))
    text_parts = []
    with fitz.open(stream=content, filetype="pdf") as document:
        logger.info("PDF document opened with %s pages", document.page_count)
        for page_number, page in enumerate(document, 1):
            page_text = page.get_text().strip()
            if page_text:
                text_parts.append(page_text)
                continue

            logger.info("Running OCR on scanned PDF page %s", page_number)
            ocr_text = _ocr_pdf_page(page, dpi).strip()
            if ocr_text:
                text_parts.append(ocr_text)
    return "\n\n".join(text_parts)


def _parse_text(content):
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("Text document could not be decoded")


def _parse_word(content):
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("Word document support is unavailable. Install 'python-docx'.") from error

    try:
        document = Document(io.BytesIO(content))
    except Exception as error:
        raise ValueError("The uploaded Word document is invalid or corrupted") from error

    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def parse_document(content, filename=None, content_type=None):
    """Extract text without an AI model.

    PDFs use selectable text when available and local OCR only for scanned pages.
    DOCX and text documents are read directly. Direct image uploads are not supported.
    """
    document_type = get_document_type(content, filename, content_type)
    if not document_type:
        raise ValueError("Unsupported file type. Upload a PDF, DOCX, TXT, MD, or CSV file.")

    logger.info("Extracting text from %s document: %s", document_type, filename or "unnamed upload")
    if document_type == "pdf":
        text = _parse_pdf(content)
    elif document_type == "word":
        text = _parse_word(content)
    else:
        text = _parse_text(content)

    logger.info("Extracted %s characters from %s document", len(text), document_type)
    return text
